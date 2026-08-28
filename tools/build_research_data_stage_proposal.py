from __future__ import annotations

import csv
import os
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
REVIEW = Path("/Users/y.k.chiang/Downloads/20260805/review_bundle_20260805")
OUT_DIR = ROOT / "output" / "reports"
OUT_DOCX = OUT_DIR / "階段研究資料整理提案_2026-08-06.docx"
CHART_DIR = OUT_DIR / "proposal_assets"

YEARS = ["2002", "2006", "2010", "2014", "2017", "2021"]

BLUE = "24557A"
DARK = "17324D"
MID = "58758C"
LIGHT = "EAF1F6"
PALE = "F5F8FA"
GOLD = "A56A00"
RED = "A23B3B"
GREEN = "2F6B55"
GRAY = "66737D"
WHITE = "FFFFFF"
BLACK = "20272C"
FONT = "Noto Sans TC"


def read_csv(relative: str) -> list[dict[str, str]]:
    with (REVIEW / relative).open("r", encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


coverage = [
    r for r in read_csv("summary_statistics/coverage_summary.csv")
    if r["sample_definition"] == "full_sample"
]
numeric = [
    r for r in read_csv("summary_statistics/numeric_summary.csv")
    if r["sample_definition"] == "full_sample"
]
categorical = [
    r for r in read_csv("summary_statistics/categorical_summary.csv")
    if r["sample_definition"] == "full_sample"
]
before_after = [
    r for r in read_csv("checks/check_income_expenditure_before_after.csv")
    if r["sample_definition"] == "full_sample"
]
sample_rows = read_csv("checks/check_analysis_sample_exclusions.csv")
age_rows = read_csv("checks/check_family_indigenous_age_counts.csv")
geo_rows = read_csv("checks/check_geography_ambiguous_or_missing.csv")
total_rows = [
    r for r in read_csv("checks/check_total_construction.csv")
    if r["sample_definition"] == "full_sample"
]

sample_idx = {r["DATA_Y"]: r for r in sample_rows}
age_idx = {r["DATA_Y"]: r for r in age_rows}
num_idx = {(r["Survey Year"], r["Variable"]): r for r in numeric}
ba_idx = {(r["survey_year"], r["variable"]): r for r in before_after}
total_idx = {(r["DATA_Y"], r["variable"]): r for r in total_rows}


def fnum(value: str | float | int | None, decimals: int = 0) -> str:
    if value is None or value == "" or str(value).upper() == "NA":
        return "—"
    x = float(value)
    return f"{x:,.{decimals}f}"


def pct(value: float | None, decimals: int = 1) -> str:
    if value is None:
        return "—"
    return f"{value * 100:.{decimals}f}%"


def mean_after(year: str, variable: str) -> float | None:
    row = ba_idx.get((year, variable))
    if not row or row["mean_after"].upper() == "NA":
        return None
    return float(row["mean_after"])


def numeric_mean(year: str, variable: str) -> float | None:
    row = num_idx.get((year, variable))
    if not row or row["Mean"].upper() == "NA":
        return None
    return float(row["Mean"])


def cat_share(year: str, variable: str, category: str) -> float | None:
    for row in categorical:
        if row["Survey Year"] == year and row["Variable"] == variable and row["Category"] == category:
            return float(row["Percentage"])
    return None


def present_vars(year: str) -> list[str]:
    return sorted(
        r["Variable"] for r in coverage
        if r["Survey Year"] == year and r["Present"] == "1"
    )


def classify_vars(year: str) -> tuple[list[str], list[str], list[str]]:
    vars_ = present_vars(year)
    core = [v for v in vars_ if not v.startswith(("INC_", "EXP_"))]
    income = [v for v in vars_ if v.startswith("INC_")]
    expenditure = [v for v in vars_ if v.startswith("EXP_")]
    return core, income, expenditure


QUALITY_LIMITED = {
    ("2002", "INC_FAM_OTHER_INCOME"),
    ("2002", "INC_FAM_WORK_INCOME"),
    ("2002", "INC_FAM_TOTAL_DERIVED_INCOME"),
    ("2002", "INC_FAM_TOTAL_DERIVED_COMPLETE_INCOME"),
    ("2006", "INC_FAM_WORK_INCOME"),
    ("2014", "RENT"),
    ("2017", "EXP_CLOTHING_EXPENDITURE"),
    ("2017", "EXP_FURNITURE_EXPENDITURE"),
    ("2017", "EXP_OTHER_EXPENDITURE"),
    ("2017", "EXP_TOTAL_DERIVED_EXPENDITURE"),
    ("2017", "EXP_TOTAL_DERIVED_COMPLETE_EXPENDITURE"),
}


def availability_mark(year: str, variable: str) -> str:
    if variable not in present_vars(year):
        return "—"
    if (year, variable) in QUALITY_LIMITED:
        return "△"
    return "●"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def set_table_fixed(table, widths: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = Inches(widths[i] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[i]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_font(run, size=10.5, bold=False, color=BLACK, italic=False) -> None:
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def style_cell_text(cell, size=8.5, bold=False, color=BLACK, align=None) -> None:
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for p in cell.paragraphs:
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(1.5)
        p.paragraph_format.line_spacing = 1.08
        if align is not None:
            p.alignment = align
        for run in p.runs:
            set_font(run, size=size, bold=bold, color=color)


def add_table(doc, headers: list[str], rows: list[list[str]], widths: list[int], small=False):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_fixed(table, widths)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    prevent_row_split(hdr)
    for i, text in enumerate(headers):
        hdr.cells[i].text = text
        set_cell_shading(hdr.cells[i], BLUE)
        style_cell_text(hdr.cells[i], size=8.0 if small else 8.5, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
    for ridx, values in enumerate(rows):
        row = table.add_row()
        prevent_row_split(row)
        if ridx % 2 == 1:
            for c in row.cells:
                set_cell_shading(c, PALE)
        for i, text in enumerate(values):
            row.cells[i].text = str(text)
            align = WD_ALIGN_PARAGRAPH.CENTER if i == 0 or (len(headers) > 3 and i > 0) else WD_ALIGN_PARAGRAPH.LEFT
            style_cell_text(row.cells[i], size=7.5 if small else 8.3, align=align)
    set_table_fixed(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_heading(doc, text: str, level: int = 1) -> None:
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)


def add_para(doc, text: str, bold_prefix: str | None = None, color=BLACK, italic=False) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.line_spacing = 1.25
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_font(r1, bold=True, color=color)
        r2 = p.add_run(text[len(bold_prefix):])
        set_font(r2, color=color, italic=italic)
    else:
        r = p.add_run(text)
        set_font(r, color=color, italic=italic)


def add_bullet(doc, text: str, color=BLACK) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    set_font(r, size=10.2, color=color)


def add_callout(doc, label: str, text: str, fill=LIGHT, color=DARK) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_fixed(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.2
    r = p.add_run(f"{label}｜")
    set_font(r, bold=True, color=color)
    r = p.add_run(text)
    set_font(r, color=color)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_source(doc, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run("資料來源：" + text)
    set_font(r, size=8, color=GRAY, italic=True)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_font(run, size=8, color=GRAY)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)
    run = paragraph.add_run(" 頁")
    set_font(run, size=8, color=GRAY)


def _chart_font(size: int, bold: bool = False):
    candidates = [
        "/private/tmp/codex-fonts/NotoSansTC.ttf",
        "/System/Library/Fonts/STHeiti Medium.ttc" if bold else "/System/Library/Fonts/STHeiti Light.ttc",
        "/System/Library/Fonts/Supplemental/Arial Unicode.ttf",
    ]
    for path in candidates:
        if Path(path).exists():
            return ImageFont.truetype(path, size=size)
    return ImageFont.load_default()


def make_charts() -> tuple[Path, Path]:
    CHART_DIR.mkdir(parents=True, exist_ok=True)
    samples = [int(sample_idx[y]["full_sample_n"]) for y in YEARS]
    analysis = [int(sample_idx[y]["indigenous_analysis_sample_n"]) for y in YEARS]
    sample_chart = CHART_DIR / "annual_samples.png"
    w, h = 1500, 600
    img = Image.new("RGB", (w, h), "white")
    d = ImageDraw.Draw(img)
    font = _chart_font(24)
    small = _chart_font(20)
    left, top, right, bottom = 110, 70, 60, 90
    plot_w, plot_h = w - left - right, h - top - bottom
    max_v = 14000
    for tick in range(0, max_v + 1, 3500):
        y = top + plot_h - tick / max_v * plot_h
        d.line((left, y, w - right, y), fill="#DCE4E9", width=2)
        d.text((15, y - 12), f"{tick:,}", font=small, fill="#52626D")
    group_w = plot_w / len(YEARS)
    for i, year in enumerate(YEARS):
        cx = left + group_w * (i + 0.5)
        for value, offset, color in ((samples[i], -24, "#4C7899"), (analysis[i], 24, "#93B5C6")):
            bar_h = value / max_v * plot_h
            d.rectangle((cx + offset - 20, top + plot_h - bar_h, cx + offset + 20, top + plot_h), fill=color)
        d.text((cx - 34, h - 70), year, font=font, fill="#283842")
    d.rectangle((1030, 20, 1060, 45), fill="#4C7899")
    d.text((1070, 17), "全樣本", font=small, fill="#283842")
    d.rectangle((1220, 20, 1250, 45), fill="#93B5C6")
    d.text((1260, 17), "原民分析樣本", font=small, fill="#283842")
    img.save(sample_chart)

    med = [mean_after(y, "EXP_MEDICAL_EXPENDITURE") for y in YEARS]
    med_chart = CHART_DIR / "medical_expenditure.png"
    img = Image.new("RGB", (w, h), "white")
    d = ImageDraw.Draw(img)
    min_v, max_v = 1700, 3000
    for tick in (1800, 2100, 2400, 2700, 3000):
        y = top + plot_h - (tick - min_v) / (max_v - min_v) * plot_h
        d.line((left, y, w - right, y), fill="#DCE4E9", width=2)
        d.text((20, y - 12), f"{tick:,}", font=small, fill="#52626D")
    points = []
    for i, (year, value) in enumerate(zip(YEARS, med)):
        x = left + plot_w * i / (len(YEARS) - 1)
        y = top + plot_h - (value - min_v) / (max_v - min_v) * plot_h
        points.append((x, y))
    d.line(points, fill="#24557A", width=6)
    for (x, y), year, value in zip(points, YEARS, med):
        d.ellipse((x - 9, y - 9, x + 9, y + 9), fill="#24557A")
        label = f"{value:,.0f}"
        bbox = d.textbbox((0, 0), label, font=small)
        d.text((x - (bbox[2] - bbox[0]) / 2, y - 40), label, font=small, fill="#24557A")
        d.text((x - 30, h - 70), year, font=font, fill="#283842")
    d.text((left, 18), "每戶每月平均醫療保健支出（元）", font=font, fill="#283842")
    img.save(med_chart)
    return sample_chart, med_chart


def build_document() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    sample_chart, med_chart = make_charts()
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.82)
    section.right_margin = Inches(0.82)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.25
    for level, size, before, after, color in ((1, 16, 16, 8, BLUE), (2, 13, 12, 6, BLUE), (3, 11.5, 9, 4, DARK)):
        st = styles[f"Heading {level}"]
        st.font.name = FONT
        st._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True
    for name in ("List Bullet", "List Number"):
        st = styles[name]
        st.font.name = FONT
        st._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        st.font.size = Pt(10.2)

    header = section.header
    hp = header.paragraphs[0]
    hp.text = "原住民族經濟狀況調查｜文健站政策研究資料整理"
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in hp.runs:
        set_font(run, size=8, color=MID)
    add_page_number(section.footer.paragraphs[0])

    # Cover: editorial cover, resolved from narrative_proposal.
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(72)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("階段研究資料整理提案")
    set_font(r, size=27, bold=True, color=DARK)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(22)
    r = p.add_run("跨年度問卷資料盤點、瑕疵評估、調整規模與文健站分析銜接")
    set_font(r, size=14, color=BLUE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(48)
    r = p.add_run("資料基準：2026-08-05 離線完整管線 review bundle\n提案日期：2026-08-06")
    set_font(r, size=10, color=GRAY)
    add_callout(
        doc,
        "核心判斷",
        "目前已可在離線受限環境建立文健站政策分析資料集。2026-08-06 v2 離線包已納入收支修正、525 站清冊、122 鄉鎮年度 panel 及個體層 exposure merge 程序；正式估計前仍須以六波個體資料重跑，確認地理匹配率與 treatment 支持範圍。",
        fill=LIGHT,
        color=DARK,
    )
    doc.add_page_break()

    add_heading(doc, "執行摘要", 1)
    add_para(doc, "本提案以問卷辨識、crosswalk、raw 匯入、主題資料建置、跨年 combined、所得支出 recode 與最終 check 檔的完整資料血緣為基礎，整理六個調查年度的可用資料、已知瑕疵及後續文健站政策分析所需工作。")
    add_bullet(doc, "六波全樣本合計 40,078 戶；各年樣本為 13,113、6,013、5,027、5,214、5,302、5,409 戶。")
    add_bullet(doc, "2006–2021 均可辨識家中是否有原住民 55+／65+ 成員；2002 問卷沒有可比家庭年齡格。")
    add_bullet(doc, "醫療保健支出六波皆可用，適合作為文健站研究的候選結果，但不能稱為就醫次數、醫療利用或健康改善。")
    add_bullet(doc, "2010 家庭其他收入、2017 三項支出 indicator、2006 未貼標籤級距代碼及 2014 RENT 解析規則已寫入離線流程；正式數值待受限環境重跑驗收。")
    add_bullet(doc, "文健站 panel 已涵蓋民國 90–114 年並含 110 年；離線流程會建立站數、首設年、事件時間及 1／3／5 年暴露，再以鄉鎮年度鍵合併個體資料。")

    add_heading(doc, "建議決策", 2)
    decision_rows = [
        ["離線重跑", "套用 v2 完整包，重建收支、文健站 panel、combined 與 exposure merge", "必要"],
        ["主分析", "建立 2006–2017 鄉鎮暴露、55+ 家庭主樣本及醫療支出分析資料集", "Go"],
        ["敏感度", "使用固定共同項所得／支出；2017 三項支出另報修正前後差異", "平行"],
        ["正式估計前", "通過六波 N、ID×年度唯一性、地理匹配率及 treatment cohort 驗收", "必要"],
    ]
    add_table(doc, ["時點", "建議", "狀態"], decision_rows, [1500, 6420, 1440])

    add_heading(doc, "一、資料處理流程與本階段資料定位", 1)
    add_para(doc, "正式管線依序執行：離線輸入檢查 → 七份 raw data 匯入 → 年度 metadata → 基本／人口／家庭資料 → 所得支出 harmonization → 跨年 combined → 描述統計與 recode 摘要 → 最終驗收。問卷 PDF 選項辨識與所得支出 crosswalk builder 是管線上游維護工具，不在每次離線重跑時重新執行。")
    flow_rows = [
        ["問卷辨識", "question_options_91_1…110", "題號、題意、選項級距", "已完成；部分表格經人工視覺修補"],
        ["Crosswalk", "unified_answer_crosswalk_*", "年度題目→整合概念", "所得與支出採概念與級距比對"],
        ["Raw 匯入", "survey_datasets.rds", "保留原始 value labels", "最新六波個體檔目前未在本機工作區"],
        ["主題資料", "basic／demographic／family／income_expenditure", "統一 ID、人口、家戶與收支", "家戶 join 六波無損失"],
        ["Combined", "cross_year_combined_data.rds", "跨主題 left join", "正式分析應以離線六波重跑版為準"],
        ["Recode／QC", "05-01、05-02、05-99", "級距中點、缺失、分布與驗收", "最新 review bundle 提供聚合證據"],
        ["政策暴露", "03-05、04-02", "建置鄉鎮年度 panel 並合併個體資料", "v2 離線包已實作；待六波重跑"],
    ]
    add_table(doc, ["階段", "主要產物", "功能", "本階段狀態"], flow_rows, [1250, 2420, 2450, 3240], small=True)

    add_heading(doc, "二、各年度資料項目列表", 1)
    add_para(doc, "下表的項目數依最新 coverage summary 中 Present=1 計算。核心／家庭欄包含人口、教育、族別、住宅、租金、家戶人數、原住民人數、55+ 結構與受訪單位；所得與支出欄包含分項及衍生總額。")
    inventory_rows = []
    for y in YEARS:
        core, inc, exp = classify_vars(y)
        inventory_rows.append([
            y,
            sample_idx[y]["full_sample_n"],
            str(len(core)),
            str(len(inc)),
            str(len(exp)),
            str(len(core) + len(inc) + len(exp)),
            "無" if y == "2002" else "有",
            "鄉鎮" if y in {"2006", "2010", "2014", "2017"} else ("部分鄉鎮" if y == "2002" else "四大區域"),
        ])
    add_table(doc, ["年度", "全樣本 N", "核心／家庭", "所得", "支出", "合計項目", "家庭55+", "地理精度"], inventory_rows, [800, 1100, 1100, 850, 850, 1000, 1000, 1760], small=True)
    add_source(doc, "coverage_summary.csv、check_family_indigenous_age_counts.csv、check_geography_ambiguous_or_missing.csv。")

    labels = {
        "2002": "核心：人口級距、教育、族別、住宅、租金、家戶／原住民人數。所得：家庭工作、政府、其他及直接回報總所得。支出：食品、醫療、交通通訊、菸酒、教育書籍合併、旅遊及直接回報總支出。",
        "2006": "新增家庭 55+／65+ 年齡結構；家庭與個人工作、政府、利息、租金、其他所得；支出擴充至照護、清潔、衣著、家具、住宅水電、貸款利息、稅保禮金等。",
        "2010": "架構近似 2006；家庭其他收入已加入問卷 I6→raw i7 例外 mapping，待離線重跑恢復整合欄。",
        "2014": "家庭／個人移轉收入加入；教育書籍合併項拆成書籍與學雜費，並加入外食住宿。家戶所得有問卷結構性適用範圍。",
        "2017": "項目架構與 2014 接近；衣著、家具家電、其他費用已依問卷指定 indicator 代碼 2 為零支出，並保留衝突稽核。",
        "2021": "項目架構與 2014／2017 接近；家庭 55+ 可用，但地理僅四大區域，無法精準連結鄉鎮文健站。",
    }
    for y in YEARS:
        add_heading(doc, y, 3)
        add_para(doc, labels[y])

    doc.add_page_break()
    add_heading(doc, "三、資料瑕疵清單與研究影響", 1)
    flaw_rows = [
        ["2017 三個支出無零值", "raw indicator 標籤亂碼", "衣著、家具家電、其他支出及總支出", "已加入 code 2=沒有；離線重跑驗收", "中"],
        ["2010 家庭其他收入全空", "問卷 I6 與 raw i7 題號差異", "INC_FAM_OTHER、2010 衍生總所得", "已加入年度例外 mapping", "小"],
        ["各年衍生總額組成不同", "問卷可得項目不同", "所得／支出總額跨波水準", "不是程式錯誤；不可當同一定義趨勢", "大"],
        ["2002 其他收入代碼斷層", "raw data 實際只使用部分代碼", "INC_FAM_OTHER、derived total", "用 reported total 或排除該分項", "中"],
        ["2002 工作所得單一級距錯配", "harmonized range mapping", "209 筆；均數影響約 +80 元", "可修，政策主模型影響極小", "小"],
        ["2006 q111 的 36 筆 1/3/4", "raw label set 與值不一致", "家庭工作所得 36/6,013", "敏感度重編；非主阻斷", "小"],
        ["2014 RENT 閉區間未解析", "字串含特殊空白", "單一租金級距", "正規化後以 2,000 元中點解析", "小"],
        ["harmonized_code 跨變數不同", "code 為變數內部代碼", "僅下游誤用時出錯", "另建 is_zero；不視為資料錯誤", "小"],
        ["C1–C4 檢查表瑕疵", "摘要口徑／命名／幽靈列", "QC 解讀，不改變個體資料", "修守門，分析可先進行", "小"],
        ["2021 無鄉鎮地理", "公開資料去識別化", "鄉鎮 treatment merge", "排除主 DID，保留區域描述", "大"],
        ["政策結果概念誤名", "舊 conceptual crosswalk", "醫療支出被稱為醫療利用", "修字典與報告用語", "中"],
        ["文健站 exposure 待正式重跑", "程序已實作、無法在本機跑六波", "處理組、事件時間與政策強度", "離線產生合併檔與匹配稽核", "中"],
        ["文健站歷史清冊存續偏誤", "由 114 年現存站回推", "過去停辦站可能漏列", "找歷年清冊或作測量誤差敏感度", "大"],
    ]
    add_table(doc, ["瑕疵／缺口", "位置／成因", "影響範圍", "研究處置", "幅度"], flaw_rows, [1850, 1900, 2000, 2890, 720], small=True)
    add_callout(doc, "判讀原則", "收支瑕疵的存在不等於文健站政策研究不可進行。只要主要結果採醫療保健支出、樣本採有原住民 55+ 成員的家戶，並排除受影響收支欄位，多數瑕疵可在平行修正中處理。", fill="EDF5F1", color=GREEN)

    add_heading(doc, "四、建議調整項目與調整幅度", 1)
    add_para(doc, "幅度同時考量程式改動、需重跑資料量、研究定義改變及驗證工作。小幅通常是單一 crosswalk／parser 或驗收表；中幅涉及一波多欄或字典語意；大幅涉及跨年度衍生變數、政策暴露或研究設計。")
    adjust_rows = [
        ["P0", "匯回最新六波個體層 combined RDS", "中", "檔案移轉、manifest/MD5、六波 N 與 key 驗證", "正式分析有可重現輸入"],
        ["P0", "離線重建文健站 exposure", "中", "民國90–114年；站數、首設年、event time、1／3／5年暴露", "panel為525站、122鄉鎮"],
        ["P0", "Treatment merge 與匹配稽核", "中", "以縣市×鄉鎮×年度合併；逐年輸出 matched/unmatched", "列數與ID×年度不變"],
        ["P0", "修正 outcome 語意", "中", "醫療支出不可稱醫療利用；確認主／次結果", "避免因果詮釋錯誤"],
        ["P1", "收支修正正式驗收", "中", "2010 mapping、2017 indicator、2006代碼、2014 RENT", "修正規則均通過離線數值檢查"],
        ["P1", "固定共同項收支版本", "小", "COMMON3 所得與 COMMON5 支出；不覆蓋 derived totals", "供跨波敏感度分析"],
        ["P1", "QC 守門重構", "中", "C1–C4、兩段式 explicit-no=0 fail、全欄變異檢查", "降低假通過"],
        ["P2", "補歷年文健站名冊", "大", "確認停辦／改名／遷站與成立年度", "降低 treatment 測量誤差"],
        ["P2", "權數盤點與加權描述統計", "中", "確認每波權數定義與校準母體", "提升母體推論"],
    ]
    add_table(doc, ["優先", "調整項目", "幅度", "工作範圍", "完成標準"], adjust_rows, [720, 2400, 850, 3250, 2140], small=True)

    add_heading(doc, "五、目前資料的敘述統計", 1)
    add_heading(doc, "5.1 年度樣本與分析樣本", 2)
    stat_rows = []
    for y in YEARS:
        a = age_idx[y]
        share55 = None if y == "2002" else int(a["has_indi_55plus_n"]) / int(a["sample_n"])
        stat_rows.append([
            y,
            fnum(sample_idx[y]["full_sample_n"]),
            fnum(sample_idx[y]["indigenous_analysis_sample_n"]),
            fnum(numeric_mean(y, "N_FAMILY"), 2),
            fnum(numeric_mean(y, "N_INDI"), 2),
            fnum(a["has_indi_55plus_n"]) if y != "2002" else "不適用",
            pct(share55),
            pct(cat_share(y, "MALE", "男性")),
        ])
    add_table(doc, ["年度", "全樣本 N", "原民分析 N", "平均家戶人數", "平均原民人數", "有原民55+戶", "占全樣本", "男性受訪者"], stat_rows, [720, 1100, 1150, 1300, 1300, 1350, 1100, 1340], small=True)
    add_source(doc, "check_analysis_sample_exclusions.csv、numeric_summary.csv、check_family_indigenous_age_counts.csv、categorical_summary.csv。男性比例受受訪單位變動影響，不宜直接解釋為母體性別趨勢。")
    doc.add_picture(str(sample_chart), width=Inches(6.15))
    cap = doc.add_paragraph("圖 1　各年度全樣本與原住民族分析樣本")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in cap.runs:
        set_font(r, size=8.5, color=GRAY)

    add_heading(doc, "5.2 所得與支出平均值", 2)
    money_rows = []
    for y in YEARS:
        money_rows.append([
            y,
            fnum(mean_after(y, "EXP_MEDICAL_EXPENDITURE")),
            fnum(mean_after(y, "INC_FAM_WORK_INCOME")),
            fnum(mean_after(y, "INC_FAM_GOV_INCOME")),
            fnum(mean_after(y, "INC_FAM_TOTAL_DERIVED_INCOME")),
            fnum(mean_after(y, "EXP_TOTAL_DERIVED_EXPENDITURE")),
        ])
    add_table(doc, ["年度", "醫療支出", "家庭工作所得", "家庭政府收入", "衍生總所得*", "衍生總支出*"], money_rows, [850, 1500, 1700, 1700, 1700, 1910], small=True)
    add_source(doc, "check_income_expenditure_before_after.csv；單位為每戶每月新臺幣元。*各年度組成項不同，僅供年度內描述，不宜直接解讀為同一定義跨年趨勢。")
    doc.add_picture(str(med_chart), width=Inches(6.15))
    cap = doc.add_paragraph("圖 2　各年度每戶每月平均醫療保健支出")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in cap.runs:
        set_font(r, size=8.5, color=GRAY)

    doc.add_page_break()
    add_heading(doc, "5.3 各年度資料樣貌", 2)
    portraits = [
        ["2002", "樣本最大；家庭規模與原民人數最高。家庭55+結構不可得；部分鄉鎮合併或缺失。醫療支出平均2,557元。"],
        ["2006", "6,013戶；51.3%家庭有原住民55+成員。受訪單位為電話接聽者，個人特徵與後期經濟戶長不可直接等同。"],
        ["2010", "5,027戶；家庭規模降至3.49人，37.0%家庭有原住民55+。家庭其他收入已加入I6→i7修正，正式值待離線重跑。"],
        ["2014", "5,214戶，全樣本中排除643戶非原住民後分析樣本4,571戶；49.2%有原住民55+。家戶所得有問卷適用範圍。"],
        ["2017", "5,302戶，原民分析樣本4,494戶；53.1%有原住民55+。三個非核心支出已加入indicator零值修正，醫療支出不受影響。"],
        ["2021", "5,409戶，原民分析樣本4,392戶；57.4%有原住民55+。地理僅四大區域，不能做鄉鎮文健站 exposure。"],
    ]
    add_table(doc, ["年度", "資料樣貌摘要"], portraits, [1000, 8360])

    doc.add_page_break()
    add_heading(doc, "六、文健站整併面向與鄉鎮成立年度分析", 1)
    add_para(doc, "2026-08-06 v2 離線包同時包含 114 年現存文健站清冊、預建鄉鎮年度 panel、panel 重建程式及個體層合併程式。清冊共 525 個現存站點、122 個鄉鎮；panel 涵蓋民國 90–114 年並補入民國 110 年。其定位是第一版政策暴露資料，不等同完整歷史營運名冊。")
    add_heading(doc, "6.1 離線整併設計", 2)
    integration_rows = [
        ["資料來源", "114年文健站營運單位清冊", "站名、縣市、行政區、成立年度；不帶入電話與地址作分析"],
        ["政策單位", "鄉鎮×民國年度", "先彙整站數再合併，避免一個鄉鎮多站造成個體重複"],
        ["個體合併鍵", "CITY×COUNTY×(DATA_Y−1911)", "CITY為縣市、COUNTY為鄉鎮；合併前統一台／臺與空白"],
        ["合併輸出", "cross_year_combined_with_care_station.rds", "保留個體層，僅能留在受限環境；預設不輸出個體CSV"],
        ["列數守門", "ID×DATA_Y唯一且合併前後列數相同", "防止站點一對多join擴張樣本"],
        ["匹配稽核", "matched／no current roster station／非鄉鎮地理", "逐調查年度輸出人數、比例及暴露概況"],
    ]
    add_table(doc, ["面向", "本版設計", "研究與資料治理考量"], integration_rows, [1700, 3050, 4610], small=True)

    add_heading(doc, "6.2 建議主分析樣本", 2)
    add_bullet(doc, "年度：2006、2010、2014、2017。2021 因無鄉鎮資料排除主模型；2002 因家庭55+不可得及部分地理模糊，作輔助前期分析。")
    add_bullet(doc, "對象：家中至少一名原住民 55+ 成員的家戶（HAS_INDI_55PLUS=TRUE）。")
    add_bullet(doc, "主要結果：家庭每月醫療保健支出；次要結果可含政府補助、照護支出或家戶經濟結果，但逐一檢查可比性。")
    add_bullet(doc, "控制變數：家戶人數、原住民人數、教育、族別、住宅等；所得控制先不放主模型，待修復後作敏感度分析。")

    add_heading(doc, "6.3 鄉鎮年度 exposure 欄位", 2)
    exposure_rows = [
        ["care_station_count", "調查年度累積站數", "連續政策強度主指標"],
        ["care_station_any", "累積站數 > 0", "二元處理狀態"],
        ["care_station_first_year", "該鄉鎮第一站成立年度", "分期設站 cohort"],
        ["care_station_new_count", "該年度新成立站數", "政策擴張流量"],
        ["care_station_event_time", "調查年－第一站成立年", "事件研究相對時間"],
        ["care_station_years_since_first", "設站後累積年數", "政策成熟度／劑量"],
        ["care_station_exposed_1y/3y/5y", "至少設站1／3／5年", "延遲效果與成熟期敏感度"],
        ["care_station_match_status", "地理匹配狀態", "界定可分析樣本與未匹配原因"],
        ["STATIONS_PER_55PLUS", "站數／55+ 原民人口或家庭", "本版尚未建立；需外部鄉鎮年度分母"],
    ]
    add_table(doc, ["欄位", "定義", "用途"], exposure_rows, [2000, 4450, 2910])

    add_heading(doc, "6.4 分析策略", 2)
    strategy_rows = [
        ["描述性", "依首次設站 cohort、站數分組比較樣本與醫療支出", "先檢查可比性與支持範圍"],
        ["事件研究", "以 FIRST_STATION_YEAR 建立相對時間", "檢驗政策前趨勢與動態效果"],
        ["分期 DID", "採適合 staggered adoption 的群組－時間 ATT", "避免傳統 TWFE 異質效果偏誤"],
        ["強度模型", "以 STATION_COUNT 或新增站數為連續暴露", "估計政策擴張與結果關聯"],
        ["異質性", "55–64／65+、山地／平地、地區、基線站數", "對應文健站服務對象與地理差異"],
        ["敏感度", "排除早期已設站、改用二元／連續 exposure、不同所得控制", "檢查估計穩健性"],
    ]
    add_table(doc, ["層次", "方法", "目的"], strategy_rows, [1500, 4550, 3310])

    add_heading(doc, "6.5 地理、時間與因果解讀限制", 2)
    limitations = [
        ["名冊存續偏誤", "114年現存清冊可能漏掉已停辦／撤站站點", "取得歷年核定與停辦清冊；以測量誤差敏感度呈現"],
        ["成立年≠完整服務年", "成立當年可能只服務部分月份", "主模型使用1年落後或exposed_1y，並比較當年起算"],
        ["早期已處理鄉鎮", "民國90年前已設站者缺乏可觀察前期", "不可作永久控制組；排除或另列早期cohort"],
        ["調查波次稀疏", "2002／2006／2010／2014／2017／2021非逐年結果", "event time需分箱；不可解讀為逐年追蹤"],
        ["2021地理不足", "僅四大區域，無法鄉鎮匹配", "排除鄉鎮主模型，保留區域描述"],
        ["2002地理與55+不足", "部分鄉鎮模糊且家庭55+不可得", "只作前期／安慰劑輔助，不作55+主模型"],
        ["服務跨界與選址", "居民可能跨鄉鎮使用；設站非隨機", "解讀為所在地政策可近性；檢查前趨勢、共變數與異質性"],
    ]
    add_table(doc, ["限制", "影響", "處理方式"], limitations, [1900, 3370, 4090], small=True)

    add_heading(doc, "6.6 正式估計前驗收條件", 2)
    checks = [
        "六波最新個體層 RDS 的 MD5、列數與 ID×DATA_Y 唯一性通過。",
        "文健站 panel 對帳為 525 站、122 鄉鎮、含民國 110 年，且各鄉鎮累積站數不下降。",
        "2006–2017 鄉鎮 exposure merge rate、unmatched 清單及行政區轉換均輸出。",
        "合併前後個體列數相同且 ID×DATA_Y 維持唯一；不得輸出或攜出個體層 CSV／RDS。",
        "每一首次設站 cohort 在處理前後至少有可用調查波次；過早處理鄉鎮不得被誤當永久控制組。",
        "醫療支出各年有效率、分布、極端級距與零值均重新確認。",
        "主樣本限定 55+ 後，逐年與 treatment cohort 的樣本數足以支持估計。",
        "標準誤至少按鄉鎮聚類；小群聚情境準備 wild cluster bootstrap。",
        "報告明確揭露 114 年現存清冊回推造成的歷史站點存續偏誤。",
    ]
    for item in checks:
        add_bullet(doc, item)

    add_heading(doc, "七、分階段執行時程", 1)
    roadmap = [
        ["第1階段", "資料凍結與基準", "匯回六波 RDS、manifest、QC 修正、研究變數清單", "1–2工作日"],
        ["第2階段", "文健站 exposure", "離線重建panel、合併個體資料、匹配稽核與cohort診斷", "1–2工作日"],
        ["第3階段", "分析資料集", "55+ 樣本、outcome、控制、merge audit、描述統計", "2–3工作日"],
        ["第4階段", "探索與主模型", "cohort 描述、event study、staggered DID、敏感度", "3–6工作日"],
        ["第5階段", "研究報告", "表圖、識別假設、限制、政策詮釋", "3–5工作日"],
    ]
    add_table(doc, ["階段", "主題", "主要產出", "預估幅度"], roadmap, [1200, 1900, 4860, 1400])

    add_heading(doc, "附錄 A　各年度可用變數對照表", 1)
    add_para(doc, "以下依 coverage summary 的 Present=1 建立跨年度對照矩陣。●＝可用；△＝欄位存在但有本提案所列品質限制；—＝該年度不可用或問卷未提供。可用只代表資料存在，不代表題意與總額組成可直接跨年比較。")

    all_present = sorted({r["Variable"] for r in coverage if r["Present"] == "1"})
    matrix_groups = [
        ("A.1 核心人口、家庭與住宅", [v for v in all_present if not v.startswith(("INC_", "EXP_"))]),
        ("A.2 所得變數", [v for v in all_present if v.startswith("INC_")]),
        ("A.3 支出變數", [v for v in all_present if v.startswith("EXP_")]),
    ]
    for title, variables in matrix_groups:
        add_heading(doc, title, 2)
        matrix_rows = [[v] + [availability_mark(y, v) for y in YEARS] for v in variables]
        count_row = ["可用項目數"] + [str(sum(v in present_vars(y) for v in variables)) for y in YEARS]
        matrix_rows.append(count_row)
        table = add_table(
            doc,
            ["整合變數"] + YEARS,
            matrix_rows,
            [3600, 960, 960, 960, 960, 960, 960],
            small=True,
        )
        for row in table.rows[1:]:
            row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            if row.cells[0].text == "可用項目數":
                for cell in row.cells:
                    set_cell_shading(cell, LIGHT)
                    style_cell_text(cell, size=7.5, bold=True, color=DARK, align=WD_ALIGN_PARAGRAPH.CENTER)
                row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_source(doc, "coverage_summary.csv；品質限制符號依本提案第三節瑕疵清單設定。")

    add_heading(doc, "附錄 B　資料來源與解讀限制", 1)
    sources = [
        "最新統計與檢查：/Users/y.k.chiang/Downloads/20260805/review_bundle_20260805/。",
        "問卷選項與 crosswalk：data/processed_data/02_metadata/question_options/、data/processed_data/03_crosswalks/。",
        "正式處理程式：code/00-00-run-remote-pipeline.R 至 code/05-99-validate-offline-pipeline.R。",
        "文健站名冊與 panel：data/raw_data/114年文健站營運單位清冊.csv、data/processed_data/05_reference/care_station_town_year_panel_long.csv。",
        "離線整併程式：code/03-05-build-care-station-town-year-panel.R、code/04-02-merge-care-station-exposure.R。",
        "本提案的平均值為未加權聚合結果；各波權數定義尚未完成跨年盤點。",
        "目前工作區的 cross_year_combined_data.rds 是本機 2021 單波版本；正式分析須改用離線 2026-08-05 六波重跑後的個體層檔案。",
    ]
    for item in sources:
        add_bullet(doc, item)

    # Keep tables with rows but allow natural pagination.
    doc.core_properties.title = "階段研究資料整理提案"
    doc.core_properties.subject = "原住民族經濟狀況調查與文健站政策分析資料準備"
    doc.core_properties.author = "Codex"
    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    build_document()
